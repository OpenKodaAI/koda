"""Tests for provider-neutral artifact ingestion."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from koda.services.artifact_ingestion import (
    ArtifactBundle,
    ArtifactKind,
    ArtifactRef,
    ArtifactStatus,
    _cache_path,
    _maybe_cleanup_artifact_cache,
    build_local_artifact_bundle,
    extract_artifact,
    extract_bundle,
)


@pytest.fixture
def artifact_cache_dir(tmp_path, monkeypatch):
    cache_dir = tmp_path / "artifact_cache"
    monkeypatch.setattr("koda.services.artifact_ingestion.ARTIFACT_CACHE_DIR", cache_dir)
    return cache_dir


def _ref(path: Path, kind: ArtifactKind, *, mime_type: str = "") -> ArtifactRef:
    return ArtifactRef(
        artifact_id=f"artifact-{path.stem}",
        kind=kind,
        label=path.name,
        source_type="test",
        mime_type=mime_type,
        path=str(path),
    )


class TestArtifactBundle:
    def test_build_local_artifact_bundle_detects_documents(self, tmp_path):
        spreadsheet = tmp_path / "report.xlsx"
        spreadsheet.write_text("placeholder", encoding="utf-8")
        bundle = build_local_artifact_bundle([str(spreadsheet)], source="telegram_document")

        assert bundle.source == "telegram_document"
        assert len(bundle.refs) == 1
        assert bundle.refs[0].kind == ArtifactKind.SPREADSHEET
        assert bundle.refs[0].label == "report.xlsx"


class TestArtifactExtraction:
    def test_cache_path_is_scoped_by_metadata_context(self, tmp_path, artifact_cache_dir):
        file_path = tmp_path / "artifact.txt"
        file_path.write_text("same payload", encoding="utf-8")
        ref_a = ArtifactRef(
            artifact_id="shared",
            kind=ArtifactKind.TEXT,
            label="artifact.txt",
            source_type="test",
            path=str(file_path),
            metadata={"workspace_root": "/tmp/a", "source_scope": ["alpha"]},
        )
        ref_b = ArtifactRef(
            artifact_id="shared",
            kind=ArtifactKind.TEXT,
            label="artifact.txt",
            source_type="test",
            path=str(file_path),
            metadata={"workspace_root": "/tmp/b", "source_scope": ["beta"]},
        )

        assert _cache_path(ref_a) != _cache_path(ref_b)

    def test_artifact_cache_cleanup_prunes_expired_entries(self, artifact_cache_dir, monkeypatch):
        artifact_cache_dir.mkdir(parents=True, exist_ok=True)
        stale_file = artifact_cache_dir / "stale.json"
        stale_file.write_text("{}", encoding="utf-8")
        monkeypatch.setattr("koda.services.artifact_ingestion._artifact_cache_last_sweep", 0.0)
        monkeypatch.setattr("koda.services.artifact_ingestion.time.time", lambda: 10_000.0)

        import os

        old_mtime = 10_000.0 - (24 * 3600) - 5
        os.utime(stale_file, (old_mtime, old_mtime))

        _maybe_cleanup_artifact_cache()

        assert not stale_file.exists()

    @pytest.mark.asyncio
    async def test_extract_bundle_enriches_refs_via_artifact_engine(self, tmp_path, artifact_cache_dir, monkeypatch):
        file_path = tmp_path / "artifact.txt"
        file_path.write_text("artifact payload", encoding="utf-8")

        class _FakeClient:
            async def start(self) -> None:
                return None

            async def stop(self) -> None:
                return None

            async def put_artifact(
                self,
                *,
                path: str,
                logical_filename: str | None = None,
                object_key: str = "",
                mime_type: str = "",
                source_metadata_json: str = "",
                purpose: str = "",
            ) -> dict[str, object]:
                assert path == str(file_path)
                assert logical_filename == file_path.name
                assert object_key == ""
                assert mime_type == "text/plain"
                assert purpose == "telegram_document"
                return {
                    "artifact_id": "content-derived-id",
                    "object_key": "agent_a/deadbeef.txt",
                    "content_hash": "deadbeef",
                    "mime_type": "text/plain",
                    "metadata_json": '{"phase":"put","size_bytes":16}',
                    "upload_outcome": "persisted_object_storage",
                }

            async def get_artifact_metadata_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                assert artifact_id == "content-derived-id"
                return {
                    "artifact_id": "content-derived-id",
                    "object_key": "agent_a/deadbeef.txt",
                    "content_hash": "deadbeef",
                    "mime_type": "text/plain",
                    "metadata_json": '{"phase":"metadata","size_bytes":16}',
                }

            async def generate_evidence_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                assert artifact_id == "content-derived-id"
                return {"evidence_json": '{"excerpt":"artifact payload"}'}

            def health(self) -> dict[str, object]:
                return {
                    "ready": True,
                    "authoritative": True,
                    "production_ready": True,
                    "cutover_allowed": True,
                    "details": {
                        "capabilities": "put_artifact,metadata,evidence",
                        "storage_backing": "object_storage_postgres",
                        "object_store": "ready",
                    },
                }

        monkeypatch.setattr(
            "koda.services.artifact_ingestion.build_artifact_engine_client",
            lambda agent_id=None: _FakeClient(),
        )

        dossier = await extract_bundle(
            ArtifactBundle(
                refs=[
                    ArtifactRef(
                        artifact_id="path-derived-id",
                        kind=ArtifactKind.TEXT,
                        label="artifact.txt",
                        source_type="test",
                        path=str(file_path),
                    )
                ],
                source="telegram_document",
            )
        )

        assert dossier.artifacts[0].ref.artifact_id == "content-derived-id"
        assert dossier.artifacts[0].ref.mime_type == "text/plain"
        assert dossier.artifacts[0].ref.metadata["content_hash"] == "deadbeef"
        assert dossier.artifacts[0].ref.metadata["object_key"] == "agent_a/deadbeef.txt"
        assert dossier.artifacts[0].ref.metadata["metadata_json"] == '{"phase":"metadata","size_bytes":16}'
        assert dossier.artifacts[0].ref.metadata["upload_outcome"] == "persisted_object_storage"
        assert dossier.artifacts[0].ref.metadata["evidence_json"] == '{"excerpt":"artifact payload"}'

    @pytest.mark.asyncio
    async def test_extract_bundle_fails_closed_when_engine_returns_empty_evidence(
        self,
        tmp_path,
        artifact_cache_dir,
        monkeypatch,
    ):
        file_path = tmp_path / "artifact.txt"
        file_path.write_text("artifact payload", encoding="utf-8")

        class _FakeClient:
            async def start(self) -> None:
                return None

            async def stop(self) -> None:
                return None

            async def put_artifact(
                self,
                *,
                path: str,
                logical_filename: str | None = None,
                object_key: str = "",
                mime_type: str = "",
                source_metadata_json: str = "",
                purpose: str = "",
            ) -> dict[str, object]:
                assert path == str(file_path)
                assert logical_filename == file_path.name
                return {
                    "artifact_id": "content-derived-id",
                    "object_key": "agent_a/deadbeef.txt",
                    "content_hash": "deadbeef",
                    "mime_type": "text/plain",
                    "metadata_json": '{"phase":"put"}',
                    "upload_outcome": "persisted_object_storage",
                }

            async def get_artifact_metadata_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                assert artifact_id == "content-derived-id"
                return {
                    "artifact_id": "content-derived-id",
                    "object_key": "agent_a/deadbeef.txt",
                    "content_hash": "deadbeef",
                    "mime_type": "text/plain",
                    "metadata_json": '{"phase":"metadata","size_bytes":16}',
                }

            async def generate_evidence_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                assert artifact_id == "content-derived-id"
                return {}

            def health(self) -> dict[str, object]:
                return {
                    "ready": True,
                    "authoritative": True,
                    "production_ready": True,
                    "cutover_allowed": True,
                    "details": {
                        "capabilities": "put_artifact,metadata,evidence",
                        "storage_backing": "object_storage_postgres",
                        "object_store": "ready",
                    },
                }

        monkeypatch.setattr(
            "koda.services.artifact_ingestion.build_artifact_engine_client",
            lambda agent_id=None: _FakeClient(),
        )

        with pytest.raises(RuntimeError, match="artifact_engine_evidence_unavailable"):
            await extract_bundle(
                ArtifactBundle(
                    refs=[
                        ArtifactRef(
                            artifact_id="path-derived-id",
                            kind=ArtifactKind.TEXT,
                            label="artifact.txt",
                            source_type="test",
                            path=str(file_path),
                        )
                    ],
                    source="telegram_document",
                )
            )

    @pytest.mark.asyncio
    async def test_extract_bundle_prefers_metadata_hash_over_ingest_hash(
        self,
        tmp_path,
        artifact_cache_dir,
        monkeypatch,
    ):
        file_path = tmp_path / "artifact.txt"
        file_path.write_text("artifact payload", encoding="utf-8")

        class _FakeClient:
            async def start(self) -> None:
                return None

            async def stop(self) -> None:
                return None

            async def put_artifact(
                self,
                *,
                path: str,
                logical_filename: str | None = None,
                object_key: str = "",
                mime_type: str = "",
                source_metadata_json: str = "",
                purpose: str = "",
            ) -> dict[str, object]:
                assert path == str(file_path)
                assert logical_filename == file_path.name
                return {
                    "artifact_id": "content-derived-id",
                    "object_key": "agent_a/deadbeef.txt",
                    "content_hash": "ingest-hash",
                    "mime_type": "text/plain",
                    "metadata_json": '{"phase":"put"}',
                    "upload_outcome": "persisted_object_storage",
                }

            async def get_artifact_metadata_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                assert artifact_id == "content-derived-id"
                return {
                    "artifact_id": "content-derived-id",
                    "object_key": "agent_a/deadbeef.txt",
                    "content_hash": "metadata-hash",
                    "mime_type": "text/plain",
                    "metadata_json": '{"phase":"metadata"}',
                }

            async def generate_evidence_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                assert artifact_id == "content-derived-id"
                return {"evidence_json": '{"excerpt":"artifact payload"}'}

            def health(self) -> dict[str, object]:
                return {
                    "ready": True,
                    "authoritative": True,
                    "production_ready": True,
                    "cutover_allowed": True,
                    "details": {
                        "capabilities": "put_artifact,metadata,evidence",
                        "storage_backing": "object_storage_postgres",
                        "object_store": "ready",
                    },
                }

        monkeypatch.setattr(
            "koda.services.artifact_ingestion.build_artifact_engine_client",
            lambda agent_id=None: _FakeClient(),
        )

        dossier = await extract_bundle(
            ArtifactBundle(
                refs=[
                    ArtifactRef(
                        artifact_id="path-derived-id",
                        kind=ArtifactKind.TEXT,
                        label="artifact.txt",
                        source_type="test",
                        path=str(file_path),
                    )
                ],
                source="telegram_document",
            )
        )

        assert dossier.artifacts[0].ref.metadata["content_hash"] == "metadata-hash"

    @pytest.mark.asyncio
    async def test_extract_bundle_fails_closed_when_ingest_fails(
        self,
        tmp_path,
        artifact_cache_dir,
        monkeypatch,
    ):
        file_path = tmp_path / "artifact.txt"
        file_path.write_text("artifact payload", encoding="utf-8")

        class _FakeClient:
            async def start(self) -> None:
                return None

            async def stop(self) -> None:
                return None

            async def put_artifact(
                self,
                *,
                path: str,
                logical_filename: str | None = None,
                object_key: str = "",
                mime_type: str = "",
                source_metadata_json: str = "",
                purpose: str = "",
            ) -> dict[str, object]:
                raise RuntimeError("put_artifact down")

            async def get_artifact_metadata_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                raise AssertionError("metadata should not be called when put_artifact fails")

            async def generate_evidence_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                raise AssertionError("evidence should not be called when put_artifact fails")

            def health(self) -> dict[str, object]:
                return {
                    "ready": True,
                    "authoritative": True,
                    "production_ready": True,
                    "cutover_allowed": True,
                    "details": {
                        "capabilities": "put_artifact,metadata,evidence",
                        "storage_backing": "object_storage_postgres",
                        "object_store": "ready",
                    },
                }

        monkeypatch.setattr(
            "koda.services.artifact_ingestion.build_artifact_engine_client",
            lambda agent_id=None: _FakeClient(),
        )

        with pytest.raises(RuntimeError, match="artifact_engine_put_unavailable"):
            await extract_bundle(
                ArtifactBundle(
                    refs=[
                        ArtifactRef(
                            artifact_id="path-derived-id",
                            kind=ArtifactKind.TEXT,
                            label="artifact.txt",
                            source_type="test",
                            path=str(file_path),
                        )
                    ],
                    source="telegram_document",
                )
            )

    @pytest.mark.asyncio
    async def test_extract_bundle_fails_closed_when_metadata_or_evidence_fail(
        self,
        tmp_path,
        artifact_cache_dir,
        monkeypatch,
    ):
        file_path = tmp_path / "artifact.txt"
        file_path.write_text("artifact payload", encoding="utf-8")

        class _FakeClient:
            async def start(self) -> None:
                return None

            async def stop(self) -> None:
                return None

            async def put_artifact(
                self,
                *,
                path: str,
                logical_filename: str | None = None,
                object_key: str = "",
                mime_type: str = "",
                source_metadata_json: str = "",
                purpose: str = "",
            ) -> dict[str, object]:
                assert path == str(file_path)
                return {
                    "artifact_id": "content-derived-id",
                    "object_key": "agent_a/deadbeef.txt",
                    "content_hash": "deadbeef",
                    "mime_type": "text/plain",
                    "metadata_json": '{"phase":"put"}',
                    "upload_outcome": "persisted_object_storage",
                }

            async def get_artifact_metadata_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                raise RuntimeError("metadata down")

            async def generate_evidence_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                raise RuntimeError("evidence down")

            def health(self) -> dict[str, object]:
                return {
                    "ready": True,
                    "authoritative": True,
                    "production_ready": True,
                    "cutover_allowed": True,
                    "details": {
                        "capabilities": "put_artifact,metadata,evidence",
                        "storage_backing": "object_storage_postgres",
                        "object_store": "ready",
                    },
                }

        monkeypatch.setattr(
            "koda.services.artifact_ingestion.build_artifact_engine_client",
            lambda agent_id=None: _FakeClient(),
        )

        with pytest.raises(RuntimeError, match="artifact_engine_metadata_unavailable"):
            await extract_bundle(
                ArtifactBundle(
                    refs=[
                        ArtifactRef(
                            artifact_id="path-derived-id",
                            kind=ArtifactKind.TEXT,
                            label="artifact.txt",
                            source_type="test",
                            path=str(file_path),
                        )
                    ],
                    source="telegram_document",
                )
            )

    @pytest.mark.asyncio
    async def test_extract_bundle_fails_closed_when_engine_reports_noncanonical_storage(
        self,
        tmp_path,
        artifact_cache_dir,
        monkeypatch,
    ):
        file_path = tmp_path / "artifact.txt"
        file_path.write_text("artifact payload", encoding="utf-8")

        class _FakeClient:
            async def start(self) -> None:
                return None

            async def stop(self) -> None:
                return None

            async def put_artifact(self, **_: object) -> dict[str, object]:
                raise AssertionError("put_artifact should not be called when storage backing is not canonical")

            async def get_artifact_metadata_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                raise AssertionError("metadata should not be called when storage backing is not canonical")

            async def generate_evidence_by_artifact_id(self, *, artifact_id: str) -> dict[str, object]:
                raise AssertionError("evidence should not be called when storage backing is not canonical")

            def health(self) -> dict[str, object]:
                return {
                    "ready": True,
                    "authoritative": True,
                    "production_ready": True,
                    "cutover_allowed": True,
                    "details": {
                        "capabilities": "put_artifact,metadata,evidence",
                        "storage_backing": "postgres_blob",
                        "object_store": "disabled",
                    },
                }

        monkeypatch.setattr(
            "koda.services.artifact_ingestion.build_artifact_engine_client",
            lambda agent_id=None: _FakeClient(),
        )

        with pytest.raises(RuntimeError, match="artifact_engine_unavailable"):
            await extract_bundle(
                ArtifactBundle(
                    refs=[
                        ArtifactRef(
                            artifact_id="path-derived-id",
                            kind=ArtifactKind.TEXT,
                            label="artifact.txt",
                            source_type="test",
                            path=str(file_path),
                        )
                    ],
                    source="telegram_document",
                )
            )

    @pytest.mark.asyncio
    @pytest.mark.parametrize(
        ("suffix", "kind", "content", "expected"),
        [
            (".txt", ArtifactKind.TEXT, "Linha 1\nLinha 2", "Linha 1 Linha 2"),
            (".json", ArtifactKind.JSON, '{"status":"ok","count":2}', '"status": "ok"'),
            (".yaml", ArtifactKind.YAML, "status: ok\ncount: 2\n", "status: ok"),
            (".xml", ArtifactKind.XML, "<root><step>Validar</step></root>", "<step> Validar"),
            (".html", ArtifactKind.HTML, "<html><body><h1>Titulo</h1><p>Resumo</p></body></html>", "Titulo Resumo"),
        ],
    )
    async def test_extract_textual_formats(self, tmp_path, artifact_cache_dir, suffix, kind, content, expected):
        file_path = tmp_path / f"artifact{suffix}"
        file_path.write_text(content, encoding="utf-8")

        extracted = await extract_artifact(_ref(file_path, kind, mime_type="text/plain"))

        assert extracted.status == ArtifactStatus.COMPLETE
        assert expected in extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_csv_and_tsv_preserve_tabular_preview(self, tmp_path, artifact_cache_dir):
        csv_path = tmp_path / "report.csv"
        csv_path.write_text("Status,Count\nOpen,3\nClosed,1\n", encoding="utf-8")
        tsv_path = tmp_path / "report.tsv"
        tsv_path.write_text("Status\tCount\nOpen\t3\nClosed\t1\n", encoding="utf-8")

        csv_extracted = await extract_artifact(_ref(csv_path, ArtifactKind.CSV, mime_type="text/csv"))
        tsv_extracted = await extract_artifact(_ref(tsv_path, ArtifactKind.TSV, mime_type="text/tab-separated-values"))

        assert csv_extracted.status == ArtifactStatus.COMPLETE
        assert "Columns: Status, Count" in csv_extracted.text_content
        assert "row 1: Open | 3" in csv_extracted.text_content
        assert tsv_extracted.status == ArtifactStatus.COMPLETE
        assert "Columns: Status, Count" in tsv_extracted.text_content
        assert "row 2: Closed | 1" in tsv_extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_docx_preserves_headings_and_tables(self, tmp_path, artifact_cache_dir):
        from docx import Document

        docx_path = tmp_path / "spec.docx"
        document = Document()
        document.add_heading("Release Notes", level=1)
        document.add_paragraph("This release fixes SIM-410.")
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Status"
        table.rows[1].cells[1].text = "Ready"
        document.save(docx_path)

        extracted = await extract_artifact(
            _ref(
                docx_path,
                ArtifactKind.DOCX,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        )

        assert extracted.status == ArtifactStatus.COMPLETE
        assert "Release Notes" in extracted.text_content
        assert "SIM-410" in extracted.text_content
        assert "Status | Ready" in extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_pdf_prefers_native_text_when_available(self, tmp_path, artifact_cache_dir):
        pdf_path = tmp_path / "native.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 fake")

        class _FakePage:
            def extract_text(self) -> str:
                return (
                    "Native PDF text with enough detail to avoid OCR fallback, "
                    "including reproduction steps and expected behavior."
                )

        class _FakeDocument:
            pages = [_FakePage()]

        with patch("pypdf.PdfReader", return_value=_FakeDocument()):
            extracted = await extract_artifact(_ref(pdf_path, ArtifactKind.PDF, mime_type="application/pdf"))

        assert extracted.status == ArtifactStatus.COMPLETE
        assert "Native PDF text" in extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_pdf_marks_partial_when_native_text_is_missing(self, tmp_path, artifact_cache_dir):
        pdf_path = tmp_path / "scan.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 fake")

        class _FakePage:
            def extract_text(self) -> str:
                return ""

        class _FakeDocument:
            pages = [_FakePage()]

        with patch("pypdf.PdfReader", return_value=_FakeDocument()):
            extracted = await extract_artifact(_ref(pdf_path, ArtifactKind.PDF, mime_type="application/pdf"))

        assert extracted.status == ArtifactStatus.PARTIAL
        assert "Could not extract readable text" in extracted.summary
        assert any("No readable text extracted from page 1." in warning for warning in extracted.warnings)

    @pytest.mark.asyncio
    async def test_extract_spreadsheet_summarizes_hidden_sheets_and_formulas(self, tmp_path, artifact_cache_dir):
        from openpyxl import Workbook

        xlsx_path = tmp_path / "metrics.xlsx"
        workbook = Workbook()
        visible = workbook.active
        visible.title = "Visible"
        visible["A1"] = "Status"
        visible["B1"] = "Count"
        visible["A2"] = "Open"
        visible["B2"] = 3
        hidden = workbook.create_sheet("Hidden")
        hidden.sheet_state = "hidden"
        hidden["A1"] = "Formula"
        hidden["B1"] = "=SUM(1,2)"
        workbook.save(xlsx_path)

        extracted = await extract_artifact(
            _ref(
                xlsx_path,
                ArtifactKind.SPREADSHEET,
                mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        )

        assert extracted.status == ArtifactStatus.COMPLETE
        assert "state=hidden" in extracted.text_content
        assert "FORMULA(=SUM(1,2))" in extracted.text_content
        assert extracted.metadata["sheet_count"] == 2

    @pytest.mark.asyncio
    async def test_extract_image_preserves_visual_path_and_ocr(self, tmp_path, artifact_cache_dir):
        from PIL import Image

        image_path = tmp_path / "screen.png"
        Image.new("RGB", (32, 32), color="white").save(image_path)

        with patch("koda.services.artifact_ingestion._ocr_image_path", return_value="Button Confirmar"):
            extracted = await extract_artifact(_ref(image_path, ArtifactKind.IMAGE, mime_type="image/png"))

        assert extracted.status == ArtifactStatus.COMPLETE
        assert extracted.visual_paths == [str(image_path)]
        assert "Button Confirmar" in extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_audio_transcribes_content(self, tmp_path, artifact_cache_dir):
        audio_path = tmp_path / "call.mp3"
        audio_path.write_bytes(b"fake-audio")

        with patch("koda.utils.audio.transcribe_audio_sync", return_value="Cliente confirmou o ajuste."):
            extracted = await extract_artifact(_ref(audio_path, ArtifactKind.AUDIO, mime_type="audio/mpeg"))

        assert extracted.status == ArtifactStatus.COMPLETE
        assert "Cliente confirmou o ajuste." in extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_video_combines_summary_frames_and_ocr(self, tmp_path, artifact_cache_dir):
        video_path = tmp_path / "flow.mp4"
        video_path.write_bytes(b"fake-video")
        frame_a = tmp_path / "frame-1.jpg"
        frame_b = tmp_path / "frame-2.jpg"
        frame_a.write_bytes(b"frame-a")
        frame_b.write_bytes(b"frame-b")

        with (
            patch(
                "koda.utils.video.process_video_attachment",
                return_value=([str(frame_a), str(frame_b)], "Video summary"),
            ),
            patch(
                "koda.services.artifact_ingestion._ocr_image_path",
                side_effect=["Tela de checkout", "Erro ao salvar"],
            ),
        ):
            extracted = await extract_artifact(_ref(video_path, ArtifactKind.VIDEO, mime_type="video/mp4"))

        assert extracted.status == ArtifactStatus.COMPLETE
        assert extracted.visual_paths == [str(frame_a), str(frame_b)]
        assert "Video summary" in extracted.text_content
        assert "Erro ao salvar" in extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_url_uses_safe_http_client(self, artifact_cache_dir):
        ref = ArtifactRef(
            artifact_id="url-1",
            kind=ArtifactKind.URL,
            label="https://example.com/runbook",
            source_type="jira_url",
            url="https://example.com/runbook",
            critical_for_action=False,
        )

        with (
            patch(
                "koda.services.http_client.inspect_url",
                return_value=MagicMock(
                    final_url="https://example.com/runbook",
                    content_type="text/html",
                    content_length=120,
                    status=200,
                ),
            ),
            patch("koda.services.http_client.fetch_url", return_value="Runbook title\nStep 1\nStep 2"),
        ):
            extracted = await extract_artifact(ref)

        assert extracted.status == ArtifactStatus.COMPLETE
        assert "Runbook title" in extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_direct_video_url_downloads_and_interprets_video(self, artifact_cache_dir, tmp_path):
        ref = ArtifactRef(
            artifact_id="url-video-1",
            kind=ArtifactKind.URL,
            label="https://cdn.example.com/demo.mp4",
            source_type="jira_url",
            url="https://cdn.example.com/demo.mp4",
        )
        frame_path = tmp_path / "frame-1.jpg"
        frame_path.write_bytes(b"frame")

        with (
            patch(
                "koda.services.http_client.inspect_url",
                side_effect=[
                    MagicMock(
                        final_url="https://cdn.example.com/demo.mp4",
                        content_type="video/mp4",
                        content_length=2048,
                        status=200,
                    ),
                    MagicMock(
                        final_url="https://cdn.example.com/demo.mp4",
                        content_type="video/mp4",
                        content_length=2048,
                        status=200,
                    ),
                ],
            ),
            patch("koda.services.http_client.download_url_bytes", return_value=b"video-bytes"),
            patch(
                "koda.utils.video.process_video_attachment",
                return_value=([str(frame_path)], "Public video summary"),
            ),
            patch("koda.services.artifact_ingestion._ocr_image_path", return_value="Tela principal"),
        ):
            extracted = await extract_artifact(ref)

        assert extracted.status == ArtifactStatus.COMPLETE
        assert extracted.ref.kind == ArtifactKind.VIDEO
        assert "Public video summary" in extracted.text_content
        assert extracted.visual_paths == [str(frame_path)]

    @pytest.mark.asyncio
    async def test_extract_video_page_uses_og_video_before_plain_html(self, artifact_cache_dir, tmp_path):
        ref = ArtifactRef(
            artifact_id="url-video-2",
            kind=ArtifactKind.URL,
            label="https://example.com/watch/demo",
            source_type="jira_url",
            url="https://example.com/watch/demo",
        )
        frame_path = tmp_path / "frame-2.jpg"
        frame_path.write_bytes(b"frame")
        html = '<html><head><meta property="og:video" content="/media/demo.mp4"></head></html>'

        with (
            patch(
                "koda.services.http_client.inspect_url",
                side_effect=[
                    MagicMock(
                        final_url="https://example.com/watch/demo",
                        content_type="text/html",
                        content_length=1024,
                        status=200,
                    ),
                    MagicMock(
                        final_url="https://example.com/media/demo.mp4",
                        content_type="video/mp4",
                        content_length=2048,
                        status=200,
                    ),
                ],
            ),
            patch("koda.services.http_client.fetch_url", return_value=html),
            patch("koda.services.http_client.download_url_bytes", return_value=b"video-bytes"),
            patch(
                "koda.utils.video.process_video_attachment",
                return_value=([str(frame_path)], "OG video summary"),
            ),
            patch("koda.services.artifact_ingestion._ocr_image_path", return_value="Erro no botão"),
        ):
            extracted = await extract_artifact(ref)

        assert extracted.status == ArtifactStatus.COMPLETE
        assert extracted.ref.kind == ArtifactKind.VIDEO
        assert "OG video summary" in extracted.text_content
        assert any("linked public video page" in warning for warning in extracted.warnings)

    @pytest.mark.asyncio
    async def test_extract_video_page_falls_back_to_ytdlp_when_needed(self, artifact_cache_dir, tmp_path):
        ref = ArtifactRef(
            artifact_id="url-video-3",
            kind=ArtifactKind.URL,
            label="https://www.youtube.com/watch?v=dQw4w9WgXcQ",
            source_type="jira_url",
            url="https://www.youtube.com/watch?v=dQw4w9WgXcQ",
        )
        frame_path = tmp_path / "frame-3.jpg"
        frame_path.write_bytes(b"frame")
        html = '<html><script type=\'application/ld+json\'>{"@type":"VideoObject"}</script></html>'

        with (
            patch(
                "koda.services.http_client.inspect_url",
                return_value=MagicMock(
                    final_url="https://www.youtube.com/watch?v=dQw4w9WgXcQ",
                    content_type="text/html",
                    content_length=512,
                    status=200,
                ),
            ),
            patch("koda.services.http_client.fetch_url", return_value=html),
            patch(
                "koda.services.artifact_ingestion._probe_video_with_ytdlp",
                return_value={"title": "Demo", "duration": 42},
            ),
            patch(
                "koda.services.artifact_ingestion._download_video_with_ytdlp",
                return_value=(b"video-bytes", "demo.mp4"),
            ),
            patch(
                "koda.utils.video.process_video_attachment",
                return_value=([str(frame_path)], "yt-dlp video summary"),
            ),
            patch("koda.services.artifact_ingestion._ocr_image_path", return_value="Frame OCR"),
        ):
            extracted = await extract_artifact(ref)

        assert extracted.status == ArtifactStatus.COMPLETE
        assert extracted.ref.kind == ArtifactKind.VIDEO
        assert "yt-dlp video summary" in extracted.text_content

    @pytest.mark.asyncio
    async def test_extract_untrusted_video_page_returns_partial_without_ytdlp(self, artifact_cache_dir):
        ref = ArtifactRef(
            artifact_id="url-video-4",
            kind=ArtifactKind.URL,
            label="https://portal.example.com/incidente/gravacao",
            source_type="jira_url",
            url="https://portal.example.com/incidente/gravacao",
        )
        html = '<html><script type=\'application/ld+json\'>{"@type":"VideoObject"}</script><p>Fluxo gravado</p></html>'

        with (
            patch(
                "koda.services.http_client.inspect_url",
                return_value=MagicMock(
                    final_url="https://portal.example.com/incidente/gravacao",
                    content_type="text/html",
                    content_length=512,
                    status=200,
                ),
            ),
            patch("koda.services.http_client.fetch_url", return_value=html),
            patch("koda.services.artifact_ingestion._probe_video_with_ytdlp") as mock_probe,
        ):
            extracted = await extract_artifact(ref)

        mock_probe.assert_not_called()
        assert extracted.status == ArtifactStatus.PARTIAL
        assert extracted.ref.kind == ArtifactKind.VIDEO
        assert "could not be fully extracted automatically" in extracted.summary
